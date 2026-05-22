#!/usr/bin/env python3
"""TechPulse-Blog 面试准备文档生成器 — 深度项目版"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── 纯灰白配色 ──────────────────────────────────────────
H1 = RGBColor(0x1A, 0x1A, 0x1A)
H2 = RGBColor(0x33, 0x33, 0x33)
H3 = RGBColor(0x55, 0x55, 0x55)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BG_TABLE_HDR = "333333"
BG_TABLE_ALT = "F8F8F8"
BG_INFO = "F5F5F5"
LINE = "E0E0E0"
BG_LIGHT = "F2F2F2"


def set_shading(el, color_hex):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    el._tc.get_or_add_tcPr().append(s)


def set_p_shading(p, color_hex):
    s = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    p._p.get_or_add_pPr().append(s)


def set_p_border_bottom(p, color="E0E0E0"):
    b = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="{color}"/></w:pBdr>')
    p._p.get_or_add_pPr().append(b)


def tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(c, BG_TABLE_HDR); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(v))
            r.font.size = Pt(9); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = BODY
            if ri % 2 == 1: set_shading(c, BG_TABLE_ALT)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Cm(w)
    sp(doc, 2)
    return t


def sp(doc, pts=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    return p


def qa(doc, num, q, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'{num}. {q}')
    r.bold = True; r.font.size = Pt(10.5); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = H1
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        if ln.startswith('##'):
            r = p.add_run(ln.replace('## ', ''))
            r.bold = True; r.font.size = Pt(10); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = H2
        elif ln.startswith('- '):
            r = p.add_run(ln)
            r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = BODY
        else:
            # Multi-line text (code blocks etc): split on \n and add line breaks
            parts = ln.split('\n')
            for i, part in enumerate(parts):
                r = p.add_run(part)
                r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = BODY
                if i < len(parts) - 1:
                    r.add_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    set_p_border_bottom(p, LINE)


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    for r in h.runs: r.font.color.rgb = H1; r.font.size = Pt(18); r.font.name = 'Microsoft YaHei'

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    for r in h.runs: r.font.color.rgb = H2; r.font.size = Pt(13); r.font.name = 'Microsoft YaHei'

def h3(doc, text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(2)
    for r in h.runs: r.font.color.rgb = H3; r.font.size = Pt(11); r.font.name = 'Microsoft YaHei'


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = BODY; r.bold = bold

def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('- ' + text)
    r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = BODY

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.name = 'Microsoft YaHei'; r.font.color.rgb = H3; r.italic = True
    set_p_shading(p, BG_INFO)


# ════════════════════════════════════════════════════════

def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    st = doc.styles['Normal']
    st.font.name = 'Microsoft YaHei'; st.font.size = Pt(10); st.font.color.rgb = BODY
    st.paragraph_format.line_spacing = 1.2

    # ── 封面 ──
    for _ in range(8): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TechPulse 技术脉动'); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = H1; r.font.name = 'Microsoft YaHei'
    sp(doc, 4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('全栈博客系统  面试准备手册'); r.font.size = Pt(13); r.font.color.rgb = H3; r.font.name = 'Microsoft YaHei'
    sp(doc, 2)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Spring Boot 2.7 / Vue 3 / Redis / RabbitMQ / AI-RAG'); r.font.size = Pt(10); r.font.color.rgb = MUTED; r.font.name = 'Microsoft YaHei'
    for _ in range(8): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('2026.05'); r.font.size = Pt(10); r.font.color.rgb = MUTED; r.font.name = 'Microsoft YaHei'
    doc.add_page_break()

    # ── 一、项目概述 ──
    h1(doc, '一、项目概述')
    body(doc, 'TechPulse 是全栈技术博客平台，集成 AI 智能助手（RAG + Function Calling），支持文章创作、社交互动、实时通知和 SSE 流式问答。')

    h2(doc, '技术栈')
    tbl(doc, ['层级', '技术选型'], [
        ['后端', 'Spring Boot 2.7 / MyBatis-Plus / MySQL 8.0 / Redis 7 / RabbitMQ 3'],
        ['前端', 'Vue 3 / Vite 5 / Tailwind CSS 4 / Pinia / STOMP.js'],
        ['AI', 'LangChain4j / DashScope Qwen-Turbo / text-embedding-v2 (RAG)'],
        ['部署', 'Docker Compose / 多阶段构建'],
    ], [2.5, 14])

    h2(doc, '面试话术')
    note(doc, '>>> 30 秒版')
    body(doc, 'TechPulse 是一个全栈博客系统，后端 Spring Boot + MyBatis-Plus，前端 Vue 3 + Tailwind CSS。技术亮点：Redis 五种使用模式——Lua 脚本原子点赞、Cache-Aside 三级防护（穿透/击穿/雪崩）、Sorted Set 热门排行、Write-Behind 计数器延迟同步；RabbitMQ 解耦通知事件 + WebSocket STOMP 实时推送；集成 LangChain4j 做 RAG 增强问答，支持 SSE 流式输出和 Function Calling。Docker Compose 一键部署。')

    note(doc, '>>> 1 分钟版')
    body(doc, '核心解决三个问题：')
    bullet(doc, '高性能缓存层：Cache-Aside + 空值缓存防穿透 + SETNX 分布式锁防击穿 + 随机 TTL 防雪崩；浏览量/点赞量 Redis 原子计数器 + CounterSyncService 定时批量同步 MySQL 最终一致；热门文章 Sorted Set ZINCRBY 实时排行。')
    bullet(doc, '异步事件驱动：点赞/评论/关注通过 RabbitMQ TopicExchange 解耦，Consumer 消费后经 WebSocket STOMP 实时推送至浏览器。通过 @ConditionalOnProperty 实现可选降级，RabbitMQ 未启动时核心业务不受影响。')
    bullet(doc, 'AI 增强问答：LangChain4j + 通义千问 RAG，文章分块(500字/100重叠)向量化后存入 InMemoryEmbeddingStore，检索 Top-5 最小相似度 0.6 注入 Prompt。SSE 流式输出(fetch+ReadableStream)、Function Calling(BlogQueryTool 4个@Tool)、每日 100 次 Redis 原子限流。')

    h2(doc, '架构')
    arch_lines = [
        '┌──────────────────────────────────────────────────┐',
        '│            浏览器 / Vue 3 前端                     │',
        '│   Axios(JWT) + STOMP.js(WS) + fetch+ReadableStream(SSE) │',
        '└───────┬──────────────┬──────────────┬─────────────┘',
        '        │ HTTP         │ WebSocket    │ SSE(POST)',
        '        ▼              ▼              ▼',
        '┌──────────────────────────────────────────────────┐',
        '│           Spring Boot 后端 (Java 17)               │',
        '│  AuthInterceptor · @RateLimit · @Idempotent · AOP  │',
        '├───────┬───────┬───────┬───────┬───────────────────┤',
        '│Controller│Service │Mapper │ Config│  Aspect/AI       │',
        '└───┬───┴───┬───┴───┬───┴───────┴───────────────────┘',
        '    ▼       ▼       ▼',
        '┌───────┐┌───────┐┌───────────┐┌──────────────────┐',
        '│MySQL 8││Redis 7││RabbitMQ 3 ││DashScope Qwen    │',
        '│9 tables││5 patterns││3 queues ││RAG+SSE+FuncCall  │',
        '└───────┘└───────┘└───────────┘└──────────────────┘',
    ]
    p = doc.add_paragraph()
    set_p_shading(p, BG_LIGHT)
    for i, line in enumerate(arch_lines):
        r = p.add_run(line)
        r.font.size = Pt(7.5); r.font.name = 'Consolas'; r.font.color.rgb = H2
        if i < len(arch_lines) - 1:
            r.add_break()
    doc.add_page_break()

    # ── 二、核心技术亮点 ──
    h1(doc, '二、核心技术亮点')

    h2(doc, '2.1 Redis 五种使用模式')
    h3(doc, '模式一：Cache-Aside 三级防护')
    bullet(doc, '穿透：空值 __NULL__ 标记，TTL 60s，防止查询不存在的数据反复穿透 DB')
    bullet(doc, '击穿：SETNX 分布式锁 + 双重检查（get→miss→tryLock→check→DB→set），锁未获取时等待 50ms 重试')
    bullet(doc, '雪崩：TTL ± 15% 随机抖动（randomJitter），避免同时过期')
    note(doc, '注意：CacheService.unlock() 未验证锁归属（未检查 value），存在误删其他请求锁的风险')

    h3(doc, '模式二：Write-Behind 原子计数器')
    bullet(doc, 'Redis INCR 原子递增 → article:dirty:{id} 脏标记（1h TTL）')
    bullet(doc, 'CounterSyncService @Scheduled 每 5 分钟 SCAN 遍历脏 key → 读 delta → UPDATE count+delta → 删 Redis')
    note(doc, '注意：sync 窗口内增量可能丢失（读delta后、删Redis前的INCR），属于最终一致性的 trade-off')

    h3(doc, '模式三：Sorted Set 热门排行')
    bullet(doc, 'hot:articles，Member=articleId，Score=浏览次数')
    bullet(doc, 'ZINCRBY +1 / ZREVRANGE Top-N，O(log(N)) 写入')
    note(doc, '注意：无时间衰减，分数只增不减；仅按浏览量排序，未加权点赞/评论')

    h3(doc, '模式四：Lua 脚本原子操作')
    bullet(doc, 'toggleSetMember(): Lua 原子执行 SISMEMBER → SREM 或 SADD + 修正 likeCount')
    bullet(doc, 'incrementWithExpire(): Lua 实现 INCR + 条件 EXPIRE（仅首次创建时设置过期）')
    bullet(doc, '为什么不用 WATCH/MULTI：乐观锁冲突率高需重试，Lua 单线程原子一次完成')

    h3(doc, '模式五：Redis 降级门面')
    bullet(doc, 'RedisUtils 每个方法 try-catch，连接异常降级到 ConcurrentHashMap')
    bullet(doc, '懒加载 TTL：get() 时检查 System.currentTimeMillis() > expireMap')
    bullet(doc, '上限 10000 条（MAX_MEMORY_CACHE_SIZE），超过拒绝写入，静默丢弃')

    h2(doc, '2.2 JWT 双 Token + 前端刷新队列')
    bullet(doc, 'Access Token(24h)：{userId, username, type:"access"}，AuthInterceptor 校验')
    bullet(doc, 'Refresh Token(7d)：{userId, type:"refresh"}，显式检查 type claim 防误用')
    bullet(doc, '前端 request.js：isRefreshing 标志 + refreshQueue 数组，401 时并发请求排队')
    bullet(doc, '刷新成功后批量重发排队请求，失败批量 reject → 跳转登录页')
    note(doc, '安全隐患：Access 和 Refresh 使用相同签名密钥；无 Token 黑名单/撤销机制')

    h2(doc, '2.3 AOP 自定义注解')
    bullet(doc, '@RateLimit：Redis Lua INCR+EXPIRE 滑动窗口，per-method-per-IP，IP 提取 X-Forwarded-For → X-Real-IP → getRemoteAddr()')
    bullet(doc, '@Idempotent：SETNX 幂等锁，Key={method}:{sha256(token)}，per-user，重复抛 409')
    note(doc, '匿名用户共享一个幂等性桶（无 userId），可能导致不同匿名用户互相阻塞')

    h2(doc, '2.4 RabbitMQ + WebSocket 通知链路')
    bullet(doc, 'Service → rabbitTemplate.convertAndSend(TopicExchange, routingKey, message)')
    bullet(doc, '3 个持久化队列：notification.like / notification.comment / notification.follow')
    bullet(doc, 'NotificationConsumer → SimpMessagingTemplate → /queue/notifications')
    bullet(doc, '@ConditionalOnProperty(name="rabbitmq.enabled", havingValue="true") 可选降级')
    note(doc, '所有 3 种通知路由到同一 WebSocket 目标，客户端需按 type 字段区分处理')

    h2(doc, '2.5 AI/RAG + SSE + Function Calling')
    bullet(doc, 'RAG 流程：文章 → 分块(500字/100重叠) → text-embedding-v2(1536维) → InMemoryEmbeddingStore')
    bullet(doc, '检索：用户问题向量化 → Top-5 最小相似度 0.6 → 注入 Prompt 上下文')
    bullet(doc, 'Function Calling：BlogQueryTool 4个 @Tool（searchArticles/getArticleSummary/getHotArticles/getRecentArticles），LLM 自主决策调用')
    bullet(doc, 'SSE 流式：后端 SseEmitter(120s 超时) + 前端 fetch() + ReadableStream 手动解析（因 EventSource 不支持 POST）')
    bullet(doc, 'ChatMemory：MessageWindowChatMemory(20轮窗口) + LRU LinkedHashMap(100会话上限, accessOrder=true)')
    bullet(doc, '限流：Redis incrementWithExpire 原子，超 100 次 → DECR 回滚 → 抛异常')
    note(doc, 'chat history 全存内存，重启丢失；generateSummary() 用同步模型，与流式模式不一致')

    h2(doc, '2.6 N+1 消除 + 前端乐观更新')
    bullet(doc, '评论列表：selectBatchIds(rootIds) 批量取子评论 → Map<rootId, List<replies>> 组装，固定 3 次 SQL')
    bullet(doc, '文章列表：selectBatchIds 批量取作者信息 → Map<userId, User> 避免 N+1')
    bullet(doc, '乐观更新：快照 prev → 立即更新 UI → API 失败 → 恢复快照 + toast.error')
    bullet(doc, 'SSE：useSSEWriter composable 封装 fetch + ReadableStream + 手动 SSE 帧解析')
    doc.add_page_break()

    # ── 三、高频面试问答（深度项目版） ──
    h1(doc, '三、高频面试问答')

    # ═══════ Redis 缓存与数据结构 ═══════
    h2(doc, 'Redis 缓存与数据结构')

    qa(doc, 1, '你的 CacheService 分布式锁实现有什么问题？', [
        '## 问题',
        'unlock() 只执行 redis.delete(lockKey)，没有验证当前锁的持有者。',
        '## 风险场景',
        '线程 A 获取锁后处理超时，锁 TTL 过期自动释放；线程 B 获取新锁；此时线程 A 完成执行调用 unlock() 会删除线程 B 的锁，导致线程 C 也能进入。',
        '## 正确做法',
        '- 加锁时 value 设为 UUID，解锁时 Lua 脚本先 GET 比较 value 再 DELETE',
        '- 或用 Redisson 的看门狗机制自动续期',
        '## 项目中的补偿',
        '- 双重检查（解锁后再次 get 缓存）降低影响',
        '- 分布式锁仅用于防击穿，即使误删也最多多查一次 DB',
    ])

    qa(doc, 2, 'incrementWithExpire 的 Lua 脚本做了什么？为什么不能直接用 INCR + EXPIRE 两条命令？', [
        '## Lua 脚本逻辑',
        'local val = redis.call("INCR", KEYS[1])\nif val == 1 then\n  redis.call("EXPIRE", KEYS[1], ARGV[1])\nend\nreturn val',
        '## 为什么需要 Lua',
        '- INCR 和 EXPIRE 是两条命令，中间有网络往返',
        '- 若 INCR 后 EXPIRE 前进程崩溃，key 永不过期（内存泄漏）',
        '- Lua 脚本在 Redis 单线程中原子执行，INCR+EXPIRE 之间不会被其他命令插入',
        '## 条件 EXPIRE 的意义',
        '- 只在 val==1（key 刚创建）时设置 TTL',
        '- 已有 TTL 的 key 不会被重置，保证过期时间稳定',
    ])

    qa(doc, 3, 'CounterSyncService Write-Behind 有什么数据一致性风险？', [
        '## 工作机制',
        'Redis INCR 原子递增 → article:dirty:{id} 脏标记(TTL 1h) → @Scheduled 每 5 分钟 SCAN → 读 delta → UPDATE count+delta → 删 Redis key',
        '## 风险：sync 窗口内增量丢失',
        '- 线程 A：SCAN 读到 delta=5，准备写 DB',
        '- 线程 B：此时 INCR +1',
        '- 线程 A：写 DB count+=5，删 Redis key',
        '- 线程 B 的 +1 丢失',
        '## 缓解措施',
        '- 脏标记 1h TTL 兜底，最多丢 5 分钟窗口数据',
        '- 非金融场景（浏览量、点赞数）最终一致可接受',
        '## 如果要完全一致',
        '- 改用 Redis + DB 双写，或 Canal 监听 binlog',
    ])

    qa(doc, 4, 'CacheService 的 randomJitter 有什么 bug？', [
        '## 代码逻辑',
        'randomJitter() 中 base = TTL_SHORT（固定值 300s），无论传入的 ttl 是多少都用 300s 作为基准计算 ±15% 抖动。',
        '## 影响',
        '- 传入 TTL_LONG(1800s) 时，抖动范围仍是 255~345s，远小于预期的 1530~2070s',
        '- 雪崩防护效果打折',
        '## 正确做法',
        '- base 应取传入的 ttl 值本身：抖动范围 = ttl * 0.85 ~ ttl * 1.15',
    ])

    qa(doc, 5, 'RedisUtils 降级到 ConcurrentHashMap 的上限是 10000，满了怎么办？', [
        '- 超过 MAX_MEMORY_CACHE_SIZE(10000) 时静默丢弃新数据',
        '- 写入操作直接 return，不抛异常，调用方感知不到',
        '- 意味着降级模式下缓存命中率会下降，更多请求穿透到 DB',
        '## 设计 trade-off',
        '- 抛异常会影响可用性，静默丢弃保可用',
        '- 10000 条对个人博客足够，高并发场景需要考虑 LRU 淘汰',
    ])

    qa(doc, 6, '你的文章点赞和评论点赞存储方式有什么不同？各有什么优缺点？', [
        '## 文章点赞（LikeServiceImpl）',
        '- Redis Set 存 user-article 关系（article:likes:{articleId}）',
        '- 只有计数器(CounterSync)同步到 MySQL，点赞关系不持久化',
        '- 优点：纯 Redis 操作极快；缺点：Redis 重启丢失所有点赞状态',
        '## 评论点赞（CommentServiceImpl）',
        '- Lua toggleSetMember 原子切换',
        '- updateCommentLikeCount() 是非原子的 read-modify-write',
        '- 存在竞态：两个并发 like 可能覆盖彼此的计数',
        '## 设计不一致',
        '- 文章点赞用 SADD/SREM，评论点赞用 Lua toggleSetMember',
        '- 应统一使用 Lua 脚本保证原子性',
    ])

    qa(doc, 7, 'Sorted Set 热门排行没有时间衰减，有什么问题？怎么改进？', [
        '## 问题',
        '- Score 只增不减，老文章浏览量累积后永远占据 Top-N',
        '- 新文章即使短期爆发也很难超越',
        '## 改进方案',
        '- 方案一：ZSet + 时间窗口，定期用 ZREMRANGEBYSCORE 清理旧数据',
        '- 方案二：分数 = 浏览量 × 时间衰减因子（如 log(1+浏览量) × e^(-λt)）',
        '- 方案三：每小时重新计算分数写入新 ZSet',
        '## 当前可接受的原因',
        '- 个人博客流量不大，排序差异不明显',
    ])

    qa(doc, 8, 'Redis 内存缓存的懒加载 TTL 有什么潜在问题？', [
        '- set() 时写入 expireMap 记录过期时间戳',
        '- get() 时检查 currentTimeMillis > expireMap.get(key)',
        '- 问题1：expireMap 本身无上限，可能内存溢出',
        '- 问题2：过期 key 不会主动清理，只在 get 时惰性删除',
        '- 问题3：多线程下 ConcurrentHashMap 但 expireMap 和 data 是两个独立 map，存在不一致窗口',
        '- 改进：用 scheduledExecutor 定期扫描清理过期 key',
    ])

    qa(doc, 9, '如果你的 Redis 集群宕机了，系统会怎样？', [
        '- RedisUtils try-catch 捕获异常 → 降级到 ConcurrentHashMap',
        '- 缓存读取全部穿透到 MySQL（空值也会回填到内存缓存）',
        '- 分布式锁失效 → 无法防击穿，热点 key 可能压垮 DB',
        '- Sorted Set 热门排行返回空（无内存降级）',
        '- @RateLimit 的 Redis INCR 失败 → 限流失效（需 fallback 逻辑）',
        '- 建议：限流应降级到 Guava RateLimiter 本地限流',
    ])

    qa(doc, 10, '你项目里 Redis 用了哪些数据类型？各用在什么场景？', [
        '- String：计数器(INCR view_count/like_count)、缓存(文章详情)、分布式锁(SETNX)',
        '- Set：点赞集合(article:likes:{id})、SISMEMBER 判断是否已赞',
        '- Sorted Set：热门排行(hot:articles)、ZINCRBY/ZREVRANGE',
        '- 未用 Hash：对象属性可用 Hash 存储，本项目用 JSON String',
        '- 未用 List：消息队列场景用 RabbitMQ 替代',
        '## 为什么不用 Hash 存文章？',
        '- 文章字段多但读取时需要全部字段，String JSON 一次获取更简单',
        '- Hash 适合部分字段更新场景',
    ])

    # ═══════ JWT 认证与安全 ═══════
    h2(doc, 'JWT 认证与安全')

    qa(doc, 11, 'Token 刷新队列是怎么处理并发 401 的？画一下流程。', [
        '## 流程',
        '1. 请求 A 收到 401，isRefreshing=true，发起 refresh 请求',
        '2. 请求 B 也收到 401，isRefreshing 已为 true，将回调推入 refreshQueue',
        '3. 请求 C 同理推入 refreshQueue',
        '4. refresh 成功 → isRefreshing=false，遍历 refreshQueue 用新 token 批量重发 A/B/C',
        '5. refresh 失败 → isRefreshing=false，遍历 refreshQueue 批量 reject → 跳登录',
        '## 关键细节',
        '- isRefreshing 是模块级变量（request.js 闭包），所有请求共享',
        '- refreshQueue 存 {resolve, reject, config}，批量 resolve 时用新 token 重发',
        '- 保证整个应用只有一个 refresh 请求在飞行',
    ])

    qa(doc, 12, 'JWT Access 和 Refresh 用同一个签名密钥有什么安全隐患？', [
        '## 风险',
        '- 如果 Access Token 泄露，攻击者可以伪造 Refresh Token（因为签名算法和密钥相同）',
        '- 理论上攻击者可以无限续签',
        '## 为什么项目这样设计',
        '- 单体应用简单场景，HS256 对称加密足够',
        '- RS256 非对称加密更安全但增加密钥管理复杂度',
        '## 改进方向',
        '- Refresh Token 用不同密钥签名',
        '- 或引入 Token 黑名单（Redis SET 存 revoked token jti）',
    ])

    qa(doc, 13, '@Idempotent 幂等注解的实现原理？匿名用户怎么处理？', [
        '## 实现',
        '- AOP @Around 拦截，Key = method + sha256(requestBody)',
        '- SETNX key "1" EX ttl',
        '- 已存在 → 抛 409 DuplicateRequestException',
        '## 匿名用户问题',
        '- 无 userId 时，Key = method:sha256(token)，所有匿名用户共享一个桶',
        '- 用户 A 的幂等 key 可能阻止用户 B 的正常请求',
        '## 改进',
        '- 匿名时用 IP + method 作为 key 维度',
        '- 或限制 @Idempotent 仅用于需登录的接口',
    ])

    qa(doc, 14, 'AuthInterceptor 的白名单怎么配置的？为什么不用 Spring Security？', [
        '## 白名单',
        '- WebMvcConfig.addInterceptors() 注册 AuthInterceptor',
        '- addPathPatterns("/api/**") + excludePathPatterns("/api/auth/login", "/api/auth/register", "/api/articles/**")',
        '## 为什么不用 Spring Security',
        '- 项目只需要 Token 校验，不需要角色权限/CSRF/记住我',
        '- Spring Security 配置复杂，学习成本高',
        '- HandlerInterceptor 更轻量可控',
        '## Spring Security 的优势',
        '- 内置攻击防护、Session 管理、Remember-Me',
        '- 适合企业级项目',
    ])

    qa(doc, 15, '如果有人拿到你的 Access Token，怎么防护？', [
        '- Access Token 24h 过期，时间窗口有限',
        '- 无黑名单机制，拿到即可用（当前缺陷）',
        '- HTTPS 传输防止中间人截获',
        '- HttpOnly Cookie 可防止 XSS 读取（项目用 Authorization Header）',
        '- 改进：短 TTL(5min) + Refresh Token + 黑名单',
    ])

    # ═══════ AOP 与 Spring ═══════
    h2(doc, 'AOP 与 Spring')

    qa(doc, 16, 'AOP 底层是 JDK 代理还是 CGLIB？你的注解用的哪种？', [
        '- Spring Boot 2.x 默认 CGLIB（基于继承，代理类是目标类子类）',
        '- JDK 代理要求目标类实现接口（基于 java.lang.reflect.Proxy）',
        '- 项目 @RateLimit/@Idempotent 标注在 Service 方法上',
        '- ServiceImpl 类未实现接口 → Spring 自动使用 CGLIB',
        '## 验证方式',
        '- debug 日志中代理类名含 $$EnhancerBySpringCGLIB$$',
    ])

    qa(doc, 17, '@Transactional 加在 deleteArticle 上，内部调用 deleteComment 事务会生效吗？', [
        '## 问题',
        '- deleteArticle() 加了 @Transactional',
        '- 如果 deleteArticle() 内部直接调用 this.deleteComment()',
        '- Spring AOP 代理拦截的是外部调用，内部 this 调用绕过代理 → 事务失效',
        '## 项目中的做法',
        '- deleteArticle() 和 deleteComment() 是独立的 Service 方法',
        '- Controller 分别调用 → 各自走代理 → 事务独立',
        '## 修复方式',
        '- 注入自身代理（@Lazy）或用 AopContext.currentProxy()',
    ])

    qa(doc, 18, '你的 @RateLimit 限流粒度是什么？有什么局限？', [
        '## 粒度',
        '- per-method-per-IP：同一 IP 对同一方法独立限流',
        '- Lua INCR 限流窗口，key = rate_limit:{methodName}:{ip}',
        '## 局限',
        '- 同一用户换 IP 就绕过（无用户级限流）',
        '- Redis 不可用时限流失效（无本地降级）',
        '- 分布式部署时多实例各自限流，实际容量 = 单实例限额 × 实例数',
        '## 改进',
        '- 加 userId 维度的限流',
        '- Redis 挂时降级到 Guava RateLimiter',
    ])

    qa(doc, 19, 'Spring Bean 的作用域？你的 ChatMemory 为什么用 synchronizedMap？', [
        '## 作用域',
        '- singleton（默认）：整个容器一个实例',
        '- prototype：每次 getBean 新建',
        '- request/session：Web 作用域',
        '## ChatMemory',
        '- AiServiceImpl 中 chatMemories 是 synchronizedMap 包装的 LRU LinkedHashMap',
        '- 因为 AiServiceImpl 是 singleton，多线程共享一个 chatMemories',
        '- synchronizedMap 保证 put/get/remove 原子性',
        '- 但 size() + isEmpty() 组合操作非原子（实际代码中不需要组合操作，所以安全）',
    ])

    # ═══════ 消息队列与实时通信 ═══════
    h2(doc, '消息队列与实时通信')

    qa(doc, 20, 'RabbitMQ TopicExchange 的 routing key 怎么设计的？为什么不用 DirectExchange？', [
        '## 设计',
        '- Exchange: blog.topic',
        '- Queue: notification.like / notification.comment / notification.follow',
        '- Routing key: notification.like / notification.comment / notification.follow',
        '- 1:1 映射（每个 routing key 对应一个 queue）',
        '## 为什么不用 Direct',
        '- Topic 支持通配符（notification.*），未来扩展灵活',
        '- 例如加 notification.system，消费者可订阅 notification.* 一次性接收',
        '## 当前实际',
        '- 因为 1:1 映射，Direct 也能工作',
        '- 用 Topic 是为了展示对消息路由的理解',
    ])

    qa(doc, 21, 'RabbitMQ 消费失败怎么处理？有死信队列吗？', [
        '## 当前处理',
        '- NotificationConsumer 中 try-catch 吞异常 + log.error',
        '- 无 requeue，消息直接 ACK（消费失败但不重试）',
        '## 风险',
        '- WebSocket 连接断开时通知丢失',
        '- 没有重试机制，一次性消费',
        '## 改进方向',
        '- 配置死信队列 DLX：消费失败 → requeue 3 次 → 进入死信队列',
        '- 或 Spring Retry 模板做指数退避重试',
    ])

    qa(doc, 22, 'WebSocket STOMP 和原生 WebSocket 有什么区别？为什么选 STOMP？', [
        '## 原生 WebSocket',
        '- 二进制/文本帧，无内置消息格式',
        '- 需自定义协议（JSON + type 字段）',
        '## STOMP',
        '- WebSocket 之上的文本协议，定义 CONNECT/SEND/SUBSCRIBE/MESSAGE 帧',
        '- Spring SimpMessagingTemplate 原生支持',
        '- 内置 @MessageMapping、@SendTo 注解',
        '- 支持用户级消息路由（/user/{id}/queue/notifications）',
        '## 选择原因',
        '- 减少自定义协议设计工作',
        '- Spring 生态集成好，代码简洁',
    ])

    qa(doc, 23, '@ConditionalOnProperty 控制 RabbitMQ 降级，具体怎么实现的？', [
        '- RabbitMQConfig 上标注 @ConditionalOnProperty(name="rabbitmq.enabled", havingValue="true")',
        '- application.yml 中 rabbitmq.enabled 默认 false',
        '- enabled=true → RabbitTemplate bean 注册 → Service 可注入',
        '- enabled=false → 无 RabbitTemplate bean → Service 中 @Autowired(required=false) 或 @Nullable',
        '- Service 发送前检查 rabbitTemplate != null',
        '- 核心业务（文章CRUD、评论、点赞）不依赖 MQ，只有通知功能降级',
    ])

    # ═══════ AI/RAG 大模型应用 ═══════
    h2(doc, 'AI/RAG 大模型应用')

    qa(doc, 24, 'RAG 的完整流程？分块策略为什么是 500字/100重叠？', [
        '## 流程',
        '1. 启动时 @EventListener(ApplicationReadyEvent) 批量索引已有文章',
        '2. 新文章发布 → ArticleSavedEvent → 增量索引',
        '3. 文章 → 分块(500字/100重叠) → text-embedding-v2(1536维) → InMemoryEmbeddingStore',
        '4. 用户提问 → 问题向量化 → store.findRelevant(question, 5, 0.6)',
        '5. 检索结果注入 Prompt → LLM 生成回答',
        '## 分块参数选择',
        '- 500字：太小会丢失上下文，太大检索不精确',
        '- 100字重叠：防止语义在块边界被截断',
        '- 500 对应约 250-300 token（中文），适合 LLM 上下文窗口',
    ])

    qa(doc, 25, 'Function Calling 和 RAG 在你的项目里怎么配合的？', [
        '## RAG',
        '- 检索知识库文章，注入 Prompt 作为上下文',
        '- 适合"博客里写了什么"这类知识查询',
        '## Function Calling',
        '- BlogQueryTool 4个 @Tool：searchArticles/getArticleSummary/getHotArticles/getRecentArticles',
        '- LLM 自主决定是否调用以及调用哪个',
        '- 适合"找某篇文章"、"最近热门"等操作性请求',
        '## 区别',
        '- RAG：被动注入上下文，LLM 基于已有信息生成',
        '- FC：LLM 主动调用 API 获取新数据，再基于结果生成',
        '- 项目中两者并行：RAG 提供知识背景，FC 提供精确操作',
    ])

    qa(doc, 26, 'SSE 流式输出的后端和前端分别怎么实现的？', [
        '## 后端',
        '- SseEmitter emitter = new SseEmitter(120000L)',
        '- executor.execute(() -> aiService.writeStream(prompt, sessionId, emitter))',
        '- writeStream 中逐 token emitter.send(SseEmitter.event().data(chunk))',
        '- 完成后 emitter.complete()',
        '## 前端',
        '- fetch() POST 请求（不用 EventSource，因为需要 POST 携带 JSON body）',
        '- response.body.getReader() 获取 ReadableStream',
        '- 手动解析 SSE 帧：data: xxx\\n\\n 格式',
        '- 逐块拼接到 assistant message.content',
        '## 超时处理',
        '- 后端 120s 超时，前端无超时控制（依赖后端完成）',
    ])

    qa(doc, 27, 'ChatMemory 的 LRU 淘汰机制？20轮窗口怎么工作的？', [
        '## 两层结构',
        '- 外层：LRU LinkedHashMap(accessOrder=true, maxCapacity=100)',
        '- 内层：MessageWindowChatMemory(maxMessages=20)',
        '## 工作流程',
        '- 用户发消息 → getOrDefault(sessionId, 创建新 Memory)',
        '- put(sessionId, memory) 触发 LRU 访问顺序更新',
        '- 超过 100 个会话时，最久未访问的会话被移除（eldestEntry 覆写）',
        '- MessageWindowChatMemory 保留最近 20 条消息，旧的自动丢弃',
        '## 问题',
        '- synchronizedMap + LRU 不是完全原子的',
        '- 内存存储，重启全部丢失',
    ])

    qa(doc, 28, 'InMemoryEmbeddingStore 重启丢数据怎么解决？', [
        '## 问题',
        '- 启动时 @EventListener 批量重新索引',
        '- 但索引过程中新文章可能未被索引',
        '## 改进方案',
        '- 方案一：用持久化向量数据库（Milvus/Qdrant/Weaviate）',
        '- 方案二：Redis Vector 搜索（Redis Stack）',
        '- 方案三：文件持久化（序列化到磁盘）',
        '## 当前可接受的原因',
        '- 个人博客文章量不大，重启后重新索引几秒内完成',
    ])

    # ═══════ 数据库设计与 SQL 优化 ═══════
    h2(doc, '数据库设计与 SQL 优化')

    qa(doc, 29, '你的 article 表设计了 deleted 字段但用硬删除，为什么？', [
        '## 代码事实',
        '- Article 实体有 @TableLogic Integer deleted 字段',
        '- 但 deleteArticle() 调用 articleMapper.deleteById()（物理删除）',
        '- MyBatis-Plus 的 deleteById 在有 @TableLogic 时应自动转为 UPDATE deleted=1',
        '- 实际执行的是硬删除，说明可能有配置问题或自定义了 SQL',
        '## 应有的行为',
        '- @TableLogic 应让 deleteById 自动变为 UPDATE deleted=1 WHERE id=?',
        '- 查询自动加 WHERE deleted=0',
        '## 如果确实是硬删除',
        '- 关联的 comment/tag 数据成为孤儿数据',
        '- 评论数、点赞数等计数器未同步清理',
    ])

    qa(doc, 30, '你的索引设计是怎样的？为什么加 FULLTEXT？', [
        '## 索引',
        '- idx_user_id：按作者查询',
        '- idx_category_id：按分类筛选',
        '- idx_status_create：状态+创建时间复合索引（分页排序）',
        '- FULLTEXT idx_title_content：标题+内容全文索引',
        '## FULLTEXT 作用',
        '- 支持 MATCH(title, content) AGAINST(?) 语法',
        '- 比 LIKE "%keyword%" 快得多（B+树 vs 全表扫描）',
        '- InnoDB 5.6+ 支持 FULLTEXT',
        '## 局限',
        '- 中文分词需要 ngram parser（MySQL 5.7.6+）',
        '- 小数据量下 FULLTEXT 优势不明显',
    ])

    qa(doc, 31, 'deleteComment 没有级联删除子评论，有什么问题？', [
        '## 问题',
        '- 删除父评论时，子评论 parentId 指向已删除的评论',
        '- getArticleComments 查询 parentId=parent.getId() 的子评论会返回空',
        '- 子评论成为孤儿数据，占用空间但不可见',
        '## 改进',
        '- 方案一：递归删除所有子评论（注意深度）',
        '- 方案二：软删除 + 查找时递归构建树',
        '- 方案三：设 parentId=null 让子评论升级为根评论',
    ])

    qa(doc, 32, '评论树是怎么构建的？为什么只支持两层？', [
        '## 构建流程',
        '1. 查根评论（WHERE parent_id IS NULL）',
        '2. 批量查子评论（WHERE parent_id IN (rootIds)）',
        '3. Map<rootId, List<Comment>> 组装',
        '## 只支持两层的原因',
        '- 简化查询：两次 selectBatchIds，固定 3 次 SQL',
        '- 多层需要递归查询或 CTE，性能随深度下降',
        '- 大多数博客评论最多两层（根评论+回复）',
        '## 展示效果',
        '- 根评论下显示回复，回复不再嵌套',
        '- 够用但不支持"回复的回复"缩进',
    ])

    qa(doc, 33, 'article 表有 tags 字段（CSV）和 article_tag 关联表，为什么有重复？', [
        '## 两套存储',
        '- article.tags = "Java,Spring,Redis"（CSV 字符串）',
        '- article_tag 关联表（article_id, tag_id）',
        '## 各自用途',
        '- CSV：列表页展示标签，一次查询无需 JOIN',
        '- 关联表：按标签筛选文章，支持规范化查询',
        '## syncArticleTags() 的 TOCTOU 问题',
        '- 先删旧关联 → 再插新关联',
        '- 并发时两个请求可能读到相同的旧状态',
        '- tag_name 不存在时先 INSERT 再 SELECT，两次查询间其他事务可能插入同名 tag',
    ])

    # ═══════ 前端工程化 ═══════
    h2(doc, '前端工程化')

    qa(doc, 34, '点赞的乐观更新失败回滚是怎么实现的？', [
        '## 流程',
        '1. 快照：prevLiked = article.isLiked; prevCount = article.likeCount',
        '2. 立即更新：article.isLiked = !prevLiked; article.likeCount += 1/-1',
        '3. 发 API：await toggleLike(id)',
        '4. 成功：用服务端返回的真实值覆盖',
        '5. 失败 catch：恢复快照 + toast.error("操作失败")',
        '## 关键点',
        '- 先更新 UI 再调 API，用户感知零延迟',
        '- 失败时精确恢复到原始状态',
        '- 使用的是 Vue 响应式，修改 ref 值即触发重渲染',
    ])

    qa(doc, 35, 'useSSEWriter 为什么用 fetch 而不是 EventSource？手动解析 SSE 帧有什么坑？', [
        '## 为什么不用 EventSource',
        '- EventSource 只支持 GET 请求',
        '- AI 聊天需要 POST 携带 JSON body（prompt + sessionId）',
        '- fetch() 支持 POST + 自定义 headers',
        '## 手动解析 SSE 帧',
        '- ReadableStream 读取的是原始字节，需要手动按 "\\n\\n" 分割帧',
        '- 每帧格式：event: xxx\\ndata: xxx\\n\\n',
        '- 需要处理 chunk 可能跨帧的情况（缓冲区管理）',
        '- 相比 EventSource 少了自动重连、Last-Event-ID 等功能',
    ])

    qa(doc, 36, '路由守卫怎么实现登录保护的？', [
        '- router.beforeEach 全局前置守卫',
        '- meta.requiresAuth 标记需要登录的路由',
        '- 检查 userStore.isLoggedIn（token 存在）',
        '- 未登录 → redirect 到 /login?redirect={当前路径}',
        '- 登录后自动跳回原页面',
        '- 白名单路由（/login, /register）不检查',
    ])

    qa(doc, 37, '虚拟列表/无限滚动在你的项目里怎么用的？', [
        '- 首页文章列表用分页加载（loadMore）',
        '- 评论列表也用分页',
        '- 未使用虚拟列表（数据量不大）',
        '- 如果评论数超过 100+，应该引入虚拟列表避免 DOM 过多',
        '## 首屏优化',
        '- 路由级懒加载（() => import()）',
        '- Vite 自动 code splitting',
        '- 图片 ProgressiveImage 渐进加载 + Unsplash fallback',
    ])

    # ═══════ Docker 与部署 ═══════
    h2(doc, 'Docker 与部署')

    qa(doc, 38, 'Docker 多阶段构建怎么做的？', [
        '## 阶段一（builder）',
        '- FROM maven:3.9-eclipse-temurin-17',
        '- COPY pom.xml + src',
        '- RUN mvn clean package -DskipTests',
        '## 阶段二（runtime）',
        '- FROM eclipse-temurin:17-jre',
        '- COPY --from=builder /app/target/*.jar app.jar',
        '- ENTRYPOINT ["java", "-jar", "app.jar"]',
        '## 优势',
        '- 最终镜像不含 Maven/源码，体积小',
        '- 依赖层缓存：pom.xml 不变时不重新下载依赖',
    ])

    qa(doc, 39, 'docker-compose.yml 中 depends_on + health check 怎么配合？', [
        '## 问题',
        '- depends_on 只保证容器启动顺序，不保证服务就绪',
        '- MySQL 容器启动 ≠ MySQL 服务就绪（需要几秒初始化）',
        '## 解决',
        '- 每个基础设施服务配置 healthcheck',
        '- MySQL: mysqladmin ping',
        '- Redis: redis-cli ping',
        '- RabbitMQ: rabbitmq-diagnostics check_port_connectivity',
        '- App 服务 depends_on 带 condition: service_healthy',
        '- 确保 MySQL/Redis/RabbitMQ 完全就绪后才启动应用',
    ])

    qa(doc, 40, 'JWT_SECRET 在 Docker 中怎么安全管理？', [
        '- docker-compose.yml 中环境变量：JWT_SECRET: ${JWT_SECRET}',
        '- .env 文件存储实际值，不提交 Git',
        '- 应用启动时 @Value("${JWT_SECRET}") 注入',
        '- 缺失时启动失败（@Value 默认 required）',
        '## 更好的做法',
        '- Docker Secrets（Swarm/K8s）',
        '- Vault 等密钥管理服务',
        '- 不要在 docker-compose.yml 中硬编码',
    ])

    # ═══════ 系统设计与开放题 ═══════
    h2(doc, '系统设计与开放题')

    qa(doc, 41, '你项目最大的技术挑战是什么？怎么解决的？', [
        '## 挑战一：Redis 三级缓存防护',
        '- 穿透/击穿/雪崩三种场景需要不同的防护策略',
        '- 分布式锁的 TTL 续期和所有权验证需要细致处理',
        '## 挑战二：SSE 流式输出',
        '- 后端 SseEmitter 超时管理',
        '- 前端 fetch+ReadableStream 手动解析 SSE 帧',
        '- 部分 chunk 跨帧的缓冲区处理',
        '## 挑战三：Write-Behind 最终一致性',
        '- Redis 计数器和 MySQL 之间的一致性窗口',
        '- CounterSyncService 的 race condition 分析',
    ])

    qa(doc, 42, '如果重新做这个项目，你会改什么？', [
        '- ES 替代 MySQL FULLTEXT（更强大的中文分词和相关性排序）',
        '- Redis Sentinel/Cluster 替代单实例（高可用）',
        '- 持久化向量数据库替代 InMemoryEmbeddingStore',
        '- 引入 Canal 监听 binlog 替代 Write-Behind 定时同步',
        '- 补充集成测试和压力测试',
        '- CI/CD 自动化（GitHub Actions）',
    ])

    qa(doc, 43, '项目有哪些已知的设计缺陷？', [
        '- CacheService.unlock() 未验证锁归属',
        '- CounterSyncService sync 窗口内增量可能丢失',
        '- 文章点赞关系不持久化到 MySQL',
        '- CommentServiceImpl.toggleLikeComment() 非原子 read-modify-write',
        '- deleteComment() 无级联删除',
        '- JWT Access/Refresh 用相同签名密钥',
        '- 无 Token 黑名单/撤销机制',
        '- article.tags CSV + article_tag 表双重存储',
        '- BCryptPasswordEncoder 直接 new（非 Spring Bean）',
        '- AiServiceImpl 限流回滚非原子',
    ])

    qa(doc, 44, '为什么选单体架构不用微服务？', [
        '## 选择原因',
        '- 个人项目，团队规模 1 人',
        '- 微服务增加运维复杂度（服务发现、配置中心、链路追踪）',
        '- Spring Boot 单体足够应对博客流量',
        '## 单体的优势',
        '- 部署简单（单 JAR/Docker）',
        '- 调试方便（一个进程内断点）',
        '- 事务管理简单（@Transactional 直接用）',
        '## 什么时候考虑微服务',
        '- 用户量 > 10万',
        '- 团队 > 5 人',
        '- 需要独立部署/扩缩容',
    ])

    qa(doc, 45, '你觉得 AI 在博客系统中最大的价值是什么？', [
        '- 内容理解：RAG 让 AI 能基于真实文章内容回答',
        '- 创作辅助：Function Calling 让 AI 能主动查询文章库',
        '- 个性化：ChatMemory 保持对话上下文',
        '- 降低门槛：用户不需要翻找文章，直接问 AI',
        '## 局限',
        '- Token 成本（Qwen-Turbo 按量计费）',
        '- RAG 质量取决于分块和 Embedding',
        '- AI 生成内容可能不准确（需 disclaimer）',
    ])

    qa(doc, 46, '你的系统能承受多大的并发？瓶颈在哪？', [
        '## 理论瓶颈',
        '- MySQL 单机：QPS 约 5000-10000',
        '- Redis 单机：QPS 10万+',
        '- SSE 连接：SseEmitter 基于 Servlet 线程池',
        '## 实际瓶颈',
        '- AiController 用 Executors.newCachedThreadPool()（无上限），可能线程爆炸',
        '- MySQL 无读写分离，写操作是瓶颈',
        '- RabbitMQ 单实例，消息堆积风险',
        '## 优化方向',
        '- 连接池配置（HikariCP maxPoolSize）',
        '- 读写分离（主从复制）',
        '- 线程池参数调谐',
    ])

    qa(doc, 47, '如果让你设计权限系统，你会怎么做？', [
        '- 当前：JWT Token + AuthInterceptor 校验登录状态',
        '- 缺失：无角色/权限控制（admin 普通用户权限相同）',
        '## 改进方案',
        '- RBAC 模型：User → Role → Permission',
        '- Spring Security + @PreAuthorize("hasRole(\'ADMIN\')")',
        '- 数据权限：文章只能编辑/删除自己的',
        '- 注解级别：@RequiresPermission("article:delete")',
    ])

    # ── 保存 ──
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TechPulse-面试准备手册.docx')
    doc.save(out)
    print(f'文档已生成：{out}')


if __name__ == '__main__':
    build()
